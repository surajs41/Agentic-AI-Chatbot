import io
import re
from typing import BinaryIO

from langchain_core.messages import HumanMessage, SystemMessage

from agentic_chatbot_backend import llm

RESUME_SYSTEM_PROMPT = """You are an expert resume writer and career coach.

Given a Job Description (JD) and a candidate's existing resume, create a tailored ONE-PAGE professional resume aligned to the JD.

Rules:
- Output exactly ONE page of content when formatted (concise bullets, no fluff)
- Use a clean professional structure with these sections in order:
  ## FULL NAME
  Contact: email | phone | location | LinkedIn (if available)
  ## PROFESSIONAL SUMMARY
  (2-3 lines tailored to the JD)
  ## CORE SKILLS
  (comma-separated, JD-relevant skills only)
  ## PROFESSIONAL EXPERIENCE
  (most relevant roles, 3-4 bullet points each, quantified where possible)
  ## EDUCATION
  (degrees only)
- Mirror JD keywords naturally — do not invent experience
- Keep tone professional and simple
- Do NOT add explanations — output ONLY the resume content in markdown
"""


def extract_resume_text(uploaded_file: BinaryIO, file_name: str) -> str:
    name = (file_name or "").lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)

    if name.endswith(".pdf"):
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()

    if name.endswith(".docx"):
        from docx import Document

        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()

    if name.endswith(".txt"):
        return data.decode("utf-8", errors="ignore").strip()

    raise ValueError("Unsupported file type. Upload PDF, DOCX, or TXT.")


def analyze_jd_and_resume(jd_text: str, resume_text: str) -> str:
    user_prompt = f"""Job Description:
{jd_text}

Candidate Resume:
{resume_text}

Create the tailored one-page resume now."""

    response = llm.invoke(
        [
            SystemMessage(content=RESUME_SYSTEM_PROMPT),
            HumanMessage(content=user_prompt),
        ]
    )
    return response.content.strip()


def _parse_sections(markdown_text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_title = "Header"
    current_lines: list[str] = []

    for line in markdown_text.splitlines():
        if line.startswith("## "):
            if current_lines:
                sections.append((current_title, "\n".join(current_lines).strip()))
            current_title = line.replace("## ", "").strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_title, "\n".join(current_lines).strip()))
    return sections


def build_resume_docx(resume_markdown: str) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    sections = _parse_sections(resume_markdown)
    if not sections:
        sections = [("Resume", resume_markdown)]

    for idx, (title, body) in enumerate(sections):
        if idx == 0 and title.lower() in {"header", "full name", "name"}:
            name_match = re.search(r"^#+\s*(.+)$", resume_markdown, re.MULTILINE)
            name = name_match.group(1).strip() if name_match else body.split("\n")[0].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(name.replace("#", "").strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            for line in body.split("\n"):
                if line.strip() and line.strip().lower() != name.lower():
                    cp = doc.add_paragraph(line.strip())
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cp.runs[0].font.size = Pt(9)
                    cp.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            continue

        heading = doc.add_paragraph()
        hr = heading.add_run(title.upper())
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0x2E, 0x5B, 0xA8)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(2)

        for line in body.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("- ") or line.startswith("* "):
                bp = doc.add_paragraph(line[2:].strip(), style="List Bullet")
                bp.paragraph_format.space_after = Pt(1)
            elif line.startswith("Contact:"):
                cp = doc.add_paragraph(line)
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.runs[0].font.size = Pt(9)
            else:
                doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
